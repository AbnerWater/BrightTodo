"""BrightToDo 文件导入待办解析服务"""

from __future__ import annotations

import base64
import io
import json
import re
import time
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import TYPE_CHECKING, Any, Literal

from lifetrace.routers.floating_capture import _call_vision_model_with_base64, get_llm_client
from lifetrace.schemas.agent import (
    ImportedCreatedTodo,
    ImportedTodoTask,
    ImportTodoFileResult,
    ImportTodosResponse,
)
from lifetrace.schemas.todo import TodoCreate, TodoPriority, TodoStatus
from lifetrace.services.agent_parse_service import LOCAL_TZ, AgentParseService
from lifetrace.util.logging_config import get_logger

if TYPE_CHECKING:
    from collections.abc import Iterable, Sequence

    from lifetrace.services.todo_service import TodoService

logger = get_logger()

MAX_IMPORT_FILE_COUNT = 5
MAX_IMPORT_FILE_BYTES = 10 * 1024 * 1024
MAX_CANDIDATE_COUNT = 40
MAX_LLM_DOCUMENT_CHARS = 8000
MAX_LLM_TODO_COUNT = 12
MAX_PREVIEW_CHARS = 1200
MAX_SOURCE_TEXT_CHARS = 240
MIN_TASK_TEXT_CHARS = 2

FileKind = Literal["image", "text", "pdf", "docx"]

IMAGE_MIME_TYPES = {"image/png", "image/jpeg", "image/webp"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp"}
TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".csv", ".json"}
TEXT_MIME_PREFIXES = ("text/",)
TEXT_MIME_TYPES = {"application/json", "application/x-ndjson"}
PDF_MIME_TYPES = {"application/pdf"}
DOCX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

TASK_HINT_RE = re.compile(
    r"(待办|任务|TODO|todo|完成|提交|复习|整理|准备|写|看|做|买|回复|截止|ddl|DDL|明天|后天|下周|今天|今晚|紧急)"
)
LIST_MARK_RE = re.compile(
    r"^\s*(?:[-*+•]|[0-9]+[.)、]|[（(]?[一二三四五六七八九十]+[）).、])\s*"
)
ISO_DURATION_RE = re.compile(
    r"^P(?:(?:\d+(?:\.\d+)?D)?"
    r"(?:T(?:\d+(?:\.\d+)?H)?(?:\d+(?:\.\d+)?M)?(?:\d+(?:\.\d+)?S)?)?)$"
)


@dataclass(frozen=True)
class AgentImportFile:
    """导入服务内部文件对象"""

    file_name: str
    mime_type: str | None
    content: bytes


@dataclass(frozen=True)
class _ProcessedFile:
    result: ImportTodoFileResult
    tasks: list[ImportedTodoTask]
    raw_text: str


class ImportTodosError(Exception):
    """导入接口契约错误"""

    def __init__(
        self,
        status_code: int,
        error_code: str,
        message: str,
        detail: str | None = None,
    ) -> None:
        super().__init__(message)
        self.status_code = status_code
        self.error_code = error_code
        self.message = message
        self.detail = detail


class FileParseError(Exception):
    """单个文件解析失败，不中断其他文件处理。"""

    def __init__(self, error_code: str, message: str, detail: str | None = None) -> None:
        super().__init__(message)
        self.error_code = error_code
        self.message = message
        self.detail = detail


class AgentImportService:
    """将上传文件解析为待确认任务。"""

    def __init__(
        self,
        parse_service: AgentParseService | None = None,
        document_llm_client: Any | None = None,
    ) -> None:
        self.parse_service = parse_service or AgentParseService()
        self._document_llm_client = document_llm_client

    def import_files(
        self,
        *,
        files: Sequence[AgentImportFile],
        reference_time: datetime | None,
        create_todos: bool,
        todo_service: TodoService,
    ) -> ImportTodosResponse:
        """解析上传文件，并按需创建草稿待办。"""
        start = time.perf_counter()
        self._validate_files(files)

        processed_files = [
            self._process_file(
                file=file,
                source_file_index=index,
                reference_time=reference_time,
                todo_service=todo_service,
            )
            for index, file in enumerate(files)
        ]
        tasks = self._dedupe_tasks(
            task for processed in processed_files for task in processed.tasks
        )
        created_todos = self._create_todos(tasks, todo_service) if create_todos else []
        preview = self._build_preview(processed.raw_text for processed in processed_files)

        return ImportTodosResponse(
            file_results=[processed.result for processed in processed_files],
            extracted_tasks=tasks,
            created_todos=created_todos,
            raw_text_preview=preview,
            processing_time_ms=round((time.perf_counter() - start) * 1000),
        )

    def _validate_files(self, files: Sequence[AgentImportFile]) -> None:
        if not files:
            raise ImportTodosError(400, "INVALID_INPUT", "请至少上传 1 个文件")
        if len(files) > MAX_IMPORT_FILE_COUNT:
            raise ImportTodosError(400, "TOO_MANY_FILES", "最多一次上传 5 个文件")

        for file in files:
            if not file.file_name:
                raise ImportTodosError(400, "INVALID_INPUT", "文件名不能为空")
            if not file.content:
                raise ImportTodosError(400, "EMPTY_FILE", f"{file.file_name} 是空文件")
            if len(file.content) > MAX_IMPORT_FILE_BYTES:
                raise ImportTodosError(
                    413,
                    "FILE_TOO_LARGE",
                    f"{file.file_name} 超过 10MB 限制",
                )
            if self._detect_file_kind(file) is None:
                raise ImportTodosError(
                    400,
                    "INVALID_FILE_TYPE",
                    "仅支持 PNG/JPEG/WebP、TXT/MD/CSV/JSON、PDF、DOCX 文件",
                    file.file_name,
                )

    def _process_file(
        self,
        *,
        file: AgentImportFile,
        source_file_index: int,
        reference_time: datetime | None,
        todo_service: TodoService,
    ) -> _ProcessedFile:
        kind = self._detect_file_kind(file)
        if kind == "image":
            return self._process_image_file(
                file,
                source_file_index,
                reference_time,
                todo_service,
            )

        try:
            raw_text = self._extract_text(file, kind)
        except FileParseError as exc:
            return _ProcessedFile(
                result=ImportTodoFileResult(
                    file_name=file.file_name,
                    mime_type=file.mime_type,
                    size_bytes=len(file.content),
                    status="failed",
                    extracted_count=0,
                    error_code=exc.error_code,
                    message=exc.message,
                    raw_text_preview="",
                ),
                tasks=[],
                raw_text="",
            )
        tasks = self._extract_tasks_from_text(
            raw_text,
            source_file=file.file_name,
            source_file_index=source_file_index,
            reference_time=reference_time,
        )
        status: Literal["success", "failed"] = "success" if raw_text.strip() else "failed"
        result = ImportTodoFileResult(
            file_name=file.file_name,
            mime_type=file.mime_type,
            size_bytes=len(file.content),
            status=status,
            extracted_count=len(tasks),
            error_code=None if status == "success" else "NO_TEXT_EXTRACTED",
            message=f"提取到 {len(tasks)} 个待确认任务" if tasks else "未识别到待办任务",
            raw_text_preview=self._limit_preview(raw_text),
        )
        return _ProcessedFile(result=result, tasks=tasks, raw_text=raw_text)

    def _process_image_file(
        self,
        file: AgentImportFile,
        source_file_index: int,
        reference_time: datetime | None,
        todo_service: TodoService,
    ) -> _ProcessedFile:
        llm_client = get_llm_client()
        if not llm_client.is_available():
            return _ProcessedFile(
                result=ImportTodoFileResult(
                    file_name=file.file_name,
                    mime_type=file.mime_type,
                    size_bytes=len(file.content),
                    status="failed",
                    extracted_count=0,
                    error_code="LLM_UNAVAILABLE",
                    message="图片解析需要先配置可用的视觉模型",
                ),
                tasks=[],
                raw_text="",
            )

        image_base64 = self._to_image_data_url(file)
        vision_todos = _call_vision_model_with_base64(
            llm_client=llm_client,
            image_base64=image_base64,
            existing_todos=self._list_existing_todos(todo_service),
        )
        tasks = self._convert_image_tasks(
            vision_todos,
            file.file_name,
            source_file_index,
            reference_time,
        )
        return _ProcessedFile(
            result=ImportTodoFileResult(
                file_name=file.file_name,
                mime_type=file.mime_type,
                size_bytes=len(file.content),
                status="success",
                extracted_count=len(tasks),
                message=f"提取到 {len(tasks)} 个待确认任务" if tasks else "图片中未识别到待办任务",
            ),
            tasks=tasks,
            raw_text="\n".join(task.source_text for task in tasks),
        )

    def _extract_text(self, file: AgentImportFile, kind: FileKind | None) -> str:
        if kind == "pdf":
            return self._extract_pdf_text(file.content)
        if kind == "docx":
            return self._extract_docx_text(file.content)
        return self._decode_text(file.content)

    def _extract_pdf_text(self, content: bytes) -> str:
        try:
            from pypdf import PdfReader  # noqa: PLC0415
        except ImportError as exc:
            raise ImportTodosError(500, "PDF_SUPPORT_MISSING", "PDF 解析依赖未安装") from exc

        try:
            reader = PdfReader(io.BytesIO(content))
            pages = [page.extract_text() or "" for page in reader.pages]
        except Exception as exc:
            raise FileParseError(
                "PDF_PARSE_FAILED",
                "PDF 文件解析失败，请确认文件未损坏",
                str(exc),
            ) from exc
        return "\n".join(page for page in pages if page.strip())

    def _extract_docx_text(self, content: bytes) -> str:
        try:
            from docx import Document  # noqa: PLC0415
        except ImportError as exc:
            raise ImportTodosError(500, "DOCX_SUPPORT_MISSING", "DOCX 解析依赖未安装") from exc

        try:
            document = Document(io.BytesIO(content))
        except Exception as exc:
            raise FileParseError(
                "DOCX_PARSE_FAILED",
                "DOCX 文件解析失败，请确认文件未损坏",
                str(exc),
            ) from exc
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    def _decode_text(self, content: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        return content.decode("utf-8", errors="ignore")

    def _extract_tasks_from_text(
        self,
        text: str,
        *,
        source_file: str,
        source_file_index: int,
        reference_time: datetime | None,
    ) -> list[ImportedTodoTask]:
        hinted_candidates = self._candidate_lines(text, require_task_hint=True)
        if hinted_candidates:
            return self._parse_candidate_tasks(
                hinted_candidates,
                source_file=source_file,
                source_file_index=source_file_index,
                reference_time=reference_time,
            )

        llm_tasks = self._extract_tasks_with_llm(
            text,
            source_file=source_file,
            source_file_index=source_file_index,
            reference_time=reference_time,
        )
        if llm_tasks:
            return llm_tasks

        fallback_task = self._infer_document_fallback_task(
            text,
            source_file=source_file,
            source_file_index=source_file_index,
        )
        if fallback_task:
            return [fallback_task]

        return self._parse_candidate_tasks(
            self._candidate_lines(text, require_task_hint=False),
            source_file=source_file,
            source_file_index=source_file_index,
            reference_time=reference_time,
        )

    def _infer_document_fallback_task(
        self,
        text: str,
        *,
        source_file: str,
        source_file_index: int,
    ) -> ImportedTodoTask | None:
        segments = self._raw_text_segments(text)
        if not segments:
            return None

        source_text = segments[0][:MAX_SOURCE_TEXT_CHARS]
        title = self._infer_document_fallback_title(text, source_file)
        duration = self._estimate_initial_duration(title)
        return ImportedTodoTask(
            task_title=title,
            priority=TodoPriority.NONE,
            due=None,
            duration=duration,
            description=(
                "LLM 暂不可用时基于文件内容生成的初始待确认任务。\n"
                f"来源文件推断：{source_text}"
            ),
            source_file=source_file,
            source_file_index=source_file_index,
            source_text=source_text,
            confidence=0.55,
        )

    def _infer_document_fallback_title(self, text: str, source_file: str) -> str:
        normalized = text.lower()
        if re.search(r"presentation|展示|汇报|答辩", normalized) and re.search(
            r"report|报告|论文|文档", normalized
        ):
            return "准备项目展示和报告"
        if re.search(r"assignment|homework|作业|课程项目|project", normalized):
            return "完成课程作业要求"
        if re.search(r"exam|quiz|考试|测验|复习", normalized):
            return "复习并整理考试材料"
        if re.search(r"meeting|minutes|会议|纪要", normalized):
            return "整理会议纪要并跟进行动项"
        if re.search(r"contract|agreement|合同|协议", normalized):
            return "审阅文件并整理待确认事项"
        stem = Path(source_file).stem or "导入文件"
        return f"阅读并整理{stem}要点"

    def _parse_candidate_tasks(
        self,
        candidates: list[str],
        *,
        source_file: str,
        source_file_index: int,
        reference_time: datetime | None,
    ) -> list[ImportedTodoTask]:
        tasks: list[ImportedTodoTask] = []
        for candidate in candidates:
            parsed = self.parse_service.parse_task(
                text=candidate,
                reference_time=reference_time,
            )
            tasks.append(
                ImportedTodoTask(
                    task_title=parsed.task_title,
                    priority=parsed.priority,
                    due=parsed.due,
                    duration=parsed.duration or self._estimate_initial_duration(parsed.task_title),
                    description=parsed.description,
                    source_file=source_file,
                    source_file_index=source_file_index,
                    source_text=candidate[:MAX_SOURCE_TEXT_CHARS],
                    confidence=parsed.confidence,
                )
            )
        return tasks

    def _candidate_lines(self, text: str, *, require_task_hint: bool) -> list[str]:
        raw_segments = self._raw_text_segments(text)
        candidates = [segment for segment in raw_segments if self._looks_like_task(segment)]
        if require_task_hint:
            return candidates[:MAX_CANDIDATE_COUNT]
        return (candidates or raw_segments[:5])[:MAX_CANDIDATE_COUNT]

    def _raw_text_segments(self, text: str) -> list[str]:
        normalized = text.replace("\r\n", "\n").replace("\r", "\n")
        raw_segments: list[str] = []
        for line in normalized.split("\n"):
            cleaned = LIST_MARK_RE.sub("", line).strip()
            if not cleaned:
                continue
            raw_segments.extend(self._split_long_line(cleaned))
        return raw_segments

    def _split_long_line(self, line: str) -> list[str]:
        if len(line) <= MAX_SOURCE_TEXT_CHARS:
            return [line]
        return [segment.strip() for segment in re.split(r"[。；;]", line) if segment.strip()]

    def _looks_like_task(self, line: str) -> bool:
        return (
            MIN_TASK_TEXT_CHARS <= len(line) <= MAX_SOURCE_TEXT_CHARS
            and TASK_HINT_RE.search(line) is not None
        )

    def _extract_tasks_with_llm(
        self,
        text: str,
        *,
        source_file: str,
        source_file_index: int,
        reference_time: datetime | None,
    ) -> list[ImportedTodoTask]:
        if not text.strip():
            return []

        llm_client = self._get_document_llm_client()
        if not llm_client.is_available():
            return []

        messages = self._build_document_llm_messages(
            text,
            source_file=source_file,
            reference_time=reference_time,
        )
        try:
            response_text = self._call_document_llm(llm_client, messages)
        except Exception as exc:
            logger.warning(f"文件导入 LLM 待办推断失败，将使用本地规则回退: {exc}")
            return []

        candidates = self._parse_llm_task_response(response_text)
        tasks = [
            self._build_task_from_llm_candidate(
                candidate,
                source_file=source_file,
                source_file_index=source_file_index,
                reference_time=reference_time,
            )
            for candidate in candidates[:MAX_LLM_TODO_COUNT]
        ]
        return [task for task in tasks if task.task_title.strip()]

    def _get_document_llm_client(self) -> Any:
        if self._document_llm_client is not None:
            return self._document_llm_client
        return get_llm_client()

    def _call_document_llm(self, llm_client: Any, messages: list[dict[str, str]]) -> str:
        if hasattr(llm_client, "_get_client") and hasattr(llm_client, "model"):
            client = llm_client._get_client()
            response = client.chat.completions.create(
                model=llm_client.model,
                messages=messages,
                temperature=0.2,
                max_tokens=1800,
                timeout=45,
                extra_body={"enable_thinking": False},
            )
            return response.choices[0].message.content or ""
        return llm_client.chat(messages=messages, temperature=0.2, max_tokens=1800)

    def _build_document_llm_messages(
        self,
        text: str,
        *,
        source_file: str,
        reference_time: datetime | None,
    ) -> list[dict[str, str]]:
        reference = self.parse_service._normalize_reference_time(reference_time)
        document_text = self._limit_document_text(text)
        system_prompt = (
            "你是 BrightToDo 的文件导入待办分析 Agent。"
            "你的任务不是只寻找显式 TODO，而是从课程资料、作业说明、会议纪要、论文、合同、项目文档中"
            "推断用户接下来需要执行、安排、跟进或准备的具体事项。"
            "每个事项必须可执行，并给出初始执行时长估计。"
            "只输出 JSON，不要输出 Markdown。"
        )
        user_prompt = (
            f"当前时间：{reference.isoformat()}\n"
            f"来源文件：{source_file}\n\n"
            "请分析下面文件文本，返回最多 12 个待确认任务。"
            "如果文档中没有明确日期，due 使用 null；如果能根据上下文推断截止时间，使用 ISO 8601 日期时间。"
            "duration 必须使用 ISO 8601 Duration，例如 PT30M、PT1H、PT2H30M。"
            "priority 只能是 high、medium、low、none。"
            "source_text 使用最能支撑该任务的原文片段。\n\n"
            "返回格式：\n"
            '{"todos":[{"title":"任务标题","description":"说明","priority":"medium",'
            '"due":null,"duration":"PT1H","source_text":"来源片段","confidence":0.8}]}\n\n'
            f"文件文本：\n{document_text}"
        )
        return [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ]

    def _limit_document_text(self, text: str) -> str:
        normalized = re.sub(r"\n{3,}", "\n\n", text.strip())
        if len(normalized) <= MAX_LLM_DOCUMENT_CHARS:
            return normalized
        return f"{normalized[:MAX_LLM_DOCUMENT_CHARS]}\n\n[后续文本已截断]"

    def _parse_llm_task_response(self, response_text: str) -> list[dict[str, Any]]:
        try:
            data = self._load_json_from_llm_response(response_text)
        except json.JSONDecodeError as exc:
            logger.warning(f"文件导入 LLM 响应不是有效 JSON: {exc}")
            return []

        if isinstance(data, list):
            todos = data
        elif isinstance(data, dict):
            todos = data.get("todos") or data.get("tasks") or data.get("new_todos") or []
        else:
            todos = []
        return [todo for todo in todos if isinstance(todo, dict)]

    def _load_json_from_llm_response(self, response_text: str) -> Any:
        cleaned = response_text.strip()
        if cleaned.startswith("```json"):
            cleaned = cleaned[7:]
        if cleaned.startswith("```"):
            cleaned = cleaned[3:]
        if cleaned.endswith("```"):
            cleaned = cleaned[:-3]
        cleaned = cleaned.strip()

        try:
            return json.loads(cleaned)
        except json.JSONDecodeError:
            json_match = re.search(r"(\{.*\}|\[.*\])", cleaned, re.DOTALL)
            if not json_match:
                raise
            return json.loads(json_match.group(1))

    def _build_task_from_llm_candidate(
        self,
        candidate: dict[str, Any],
        *,
        source_file: str,
        source_file_index: int,
        reference_time: datetime | None,
    ) -> ImportedTodoTask:
        title = str(candidate.get("title") or candidate.get("name") or "").strip()
        source_text = str(candidate.get("source_text") or title).strip()
        parsed = self.parse_service.parse_task(
            text=source_text or title,
            reference_time=reference_time,
        )
        priority = self._normalize_priority(candidate.get("priority"), parsed.priority)
        due = self._parse_llm_due(candidate.get("due")) or parsed.due
        duration = (
            self._normalize_llm_duration(candidate.get("duration"))
            or parsed.duration
            or self._estimate_initial_duration(title or parsed.task_title)
        )
        description = self._build_llm_task_description(candidate, source_text, duration)
        confidence = self._normalize_confidence(candidate.get("confidence"), parsed.confidence)
        return ImportedTodoTask(
            task_title=title or parsed.task_title,
            priority=priority,
            due=due,
            duration=duration,
            description=description,
            source_file=source_file,
            source_file_index=source_file_index,
            source_text=(source_text or title)[:MAX_SOURCE_TEXT_CHARS],
            confidence=confidence,
        )

    def _normalize_priority(self, value: Any, fallback: TodoPriority) -> TodoPriority:
        raw = str(value or "").strip().lower()
        if raw in {"high", "medium", "low", "none"}:
            return TodoPriority(raw)
        return fallback

    def _parse_llm_due(self, value: Any) -> datetime | None:
        if value in (None, ""):
            return None
        raw = str(value).strip()
        if raw.lower() in {"none", "null", "unknown", "无"}:
            return None
        try:
            parsed = datetime.fromisoformat(raw.replace("Z", "+00:00"))
        except ValueError:
            return None
        if parsed.tzinfo is None:
            return parsed.replace(tzinfo=LOCAL_TZ)
        return parsed

    def _normalize_llm_duration(self, value: Any) -> str | None:
        if value in (None, ""):
            return None
        raw = str(value).strip().upper()
        if ISO_DURATION_RE.match(raw) and raw not in {"P", "PT"}:
            return raw

        hour_match = re.search(r"(?P<hours>\d+(?:\.\d+)?)\s*(?:H|小时|HOUR)", raw, re.I)
        minute_match = re.search(r"(?P<minutes>\d+)\s*(?:M|分钟|MIN)", raw, re.I)
        minutes = 0
        if hour_match:
            minutes += int(float(hour_match.group("hours")) * 60)
        if minute_match:
            minutes += int(minute_match.group("minutes"))
        return self._format_duration(minutes) if minutes > 0 else None

    def _estimate_initial_duration(self, title: str) -> str:
        if re.search(r"回复|确认|提交表单|发送", title):
            return "PT30M"
        if re.search(r"阅读|预习|复习|梳理|整理", title):
            return "PT1H"
        if re.search(r"报告|论文|作业|项目|实现|设计|展示|汇报", title):
            return "PT2H"
        return "PT1H"

    def _format_duration(self, minutes: int) -> str:
        if minutes % 60 == 0:
            return f"PT{minutes // 60}H"
        return f"PT{minutes}M"

    def _build_llm_task_description(
        self,
        candidate: dict[str, Any],
        source_text: str,
        duration: str | None,
    ) -> str:
        parts: list[str] = []
        description = str(candidate.get("description") or "").strip()
        if description:
            parts.append(description)
        if source_text:
            parts.append(f"来源文件推断：{source_text[:MAX_SOURCE_TEXT_CHARS]}")
        if duration:
            parts.append(f"初始时长估计：{duration}")
        return "\n".join(parts)

    def _normalize_confidence(self, value: Any, fallback: float) -> float:
        try:
            confidence = float(value)
        except (TypeError, ValueError):
            confidence = fallback
        return max(0.0, min(round(confidence, 2), 1.0))

    def _convert_image_tasks(
        self,
        todos: list[dict[str, object]],
        source_file: str,
        source_file_index: int,
        reference_time: datetime | None,
    ) -> list[ImportedTodoTask]:
        tasks: list[ImportedTodoTask] = []
        for todo in todos:
            title = str(todo.get("title") or "").strip()
            source_text = str(todo.get("source_text") or title).strip()
            if not title and not source_text:
                continue
            parsed = self.parse_service.parse_task(
                text=source_text or title,
                reference_time=reference_time,
            )
            tasks.append(
                ImportedTodoTask(
                    task_title=title or parsed.task_title,
                    priority=parsed.priority,
                    due=parsed.due,
                    duration=parsed.duration,
                    description=str(todo.get("description") or parsed.description or ""),
                    source_file=source_file,
                    source_file_index=source_file_index,
                    source_text=(source_text or title)[:MAX_SOURCE_TEXT_CHARS],
                    confidence=float(todo.get("confidence") or parsed.confidence),
                )
            )
        return tasks

    def _create_todos(
        self,
        tasks: Sequence[ImportedTodoTask],
        todo_service: TodoService,
    ) -> list[ImportedCreatedTodo]:
        created: list[ImportedCreatedTodo] = []
        try:
            for task in tasks:
                todo = todo_service.create_todo(
                    TodoCreate(
                        name=task.task_title,
                        description=task.description,
                        user_notes=self._build_user_notes(task),
                        due=task.due,
                        duration=task.duration,
                        status=TodoStatus.DRAFT,
                        priority=task.priority,
                        tags=["文件导入", "AI解析"],
                    )
                )
                created.append(
                    ImportedCreatedTodo(id=todo.id, name=todo.name, status=todo.status)
                )
        except Exception as exc:
            self._rollback_created_todos(todo_service, created)
            raise ImportTodosError(
                500,
                "TODO_CREATE_FAILED",
                "创建待办失败，已回滚本次已创建项",
                str(exc),
            ) from exc
        return created

    def _rollback_created_todos(
        self,
        todo_service: TodoService,
        created: Sequence[ImportedCreatedTodo],
    ) -> None:
        for todo in reversed(created):
            try:
                todo_service.delete_todo(todo.id)
            except Exception:
                continue

    def _build_user_notes(self, task: ImportedTodoTask) -> str:
        return (
            "来源：文件导入\n"
            f"来源文件：{task.source_file}\n"
            f"来源文本：{task.source_text}\n"
            f"置信度：{round(task.confidence * 100)}%"
        )

    def _dedupe_tasks(self, tasks: Iterable[ImportedTodoTask]) -> list[ImportedTodoTask]:
        result: list[ImportedTodoTask] = []
        seen: set[tuple[str, str | None]] = set()
        for task in tasks:
            key = (re.sub(r"\s+", "", task.task_title).lower(), task.due.isoformat() if task.due else None)
            if key in seen:
                continue
            seen.add(key)
            result.append(task)
        return result

    def _build_preview(self, text_parts: Iterable[str]) -> str:
        combined = "\n\n".join(part.strip() for part in text_parts if part.strip())
        return self._limit_preview(combined)

    def _limit_preview(self, text: str) -> str:
        text = text.strip()
        if len(text) <= MAX_PREVIEW_CHARS:
            return text
        return f"{text[:MAX_PREVIEW_CHARS]}..."

    def _detect_file_kind(self, file: AgentImportFile) -> FileKind | None:
        suffix = Path(file.file_name).suffix.lower()
        mime_type = (file.mime_type or "").lower()
        if suffix in IMAGE_EXTENSIONS or mime_type in IMAGE_MIME_TYPES:
            return "image"
        if suffix in TEXT_EXTENSIONS or mime_type.startswith(TEXT_MIME_PREFIXES) or mime_type in TEXT_MIME_TYPES:
            return "text"
        if suffix == ".pdf" or mime_type in PDF_MIME_TYPES:
            return "pdf"
        if suffix == ".docx" or mime_type in DOCX_MIME_TYPES:
            return "docx"
        return None

    def _to_image_data_url(self, file: AgentImportFile) -> str:
        mime_type = file.mime_type or self._mime_from_extension(file.file_name)
        encoded = base64.b64encode(file.content).decode("ascii")
        return f"data:{mime_type};base64,{encoded}"

    def _mime_from_extension(self, file_name: str) -> str:
        suffix = Path(file_name).suffix.lower()
        if suffix == ".webp":
            return "image/webp"
        if suffix in {".jpg", ".jpeg"}:
            return "image/jpeg"
        return "image/png"

    def _list_existing_todos(self, todo_service: TodoService) -> list[dict[str, object]]:
        payload = todo_service.list_todos(limit=1000, offset=0, status=None)
        todos = payload.get("todos", [])
        return [
            {"id": todo.id, "name": todo.name, "description": todo.description}
            for todo in todos
        ]
